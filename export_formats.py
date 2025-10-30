"""
export_formats.py
Export grant proposals to DOCX and PDF formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from typing import Dict
from datetime import datetime

class ProposalExporter:
    """
    Export grant proposals to professional document formats
    
    Supports:
    - Microsoft Word (.docx)
    - PDF (via docx2pdf or reportlab)
    - HTML
    - Markdown
    """
    
    def __init__(self):
        self.doc = None

    def export_to_docx(
        self,
        proposal_data: Dict,
        output_path: str = "grant_proposal.docx",
        template: str = "nsf"
    ) -> str:
        """
        Export to Microsoft Word format
        
        Args:
            proposal_data: Complete proposal data from Phase 3
            output_path: Path for output file
            template: Template style ('nsf', 'nih', 'simple')
        
        Returns:
            Path to created file
        """
        print(f"📝 Exporting to DOCX: {output_path}")
        
        # Create document
        self.doc = Document()
        
        # Set up styles
        self._setup_styles()
        
        # Add content based on template
        if template == "nsf":
            self._format_nsf_proposal(proposal_data)
        elif template == "nih":
            self._format_nih_proposal(proposal_data)
        else:
            self._format_simple_proposal(proposal_data)
        
        # Save
        self.doc.save(output_path)
        print(f"✅ DOCX saved: {output_path}")
        
        return output_path
    
    def _setup_styles(self):
        """Set up document styles"""
        styles = self.doc.styles
        
        # Title style
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(18)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 0, 128)
        
        # Heading 1 style
        heading1 = styles['Heading 1']
        heading1_font = heading1.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(14)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(0, 0, 0)
        
        # Normal style
        normal = styles['Normal']
        normal_font = normal.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
    
    def _format_nsf_proposal(self, data: Dict):
        """Format as NSF-style proposal"""
        
        analysis = data.get('analysis', {})
        evaluation = data.get('evaluation', {})
        innovations = data.get('innovations', {})
        proposal = data.get('proposal', {}).get('proposal', {})
        
        # Cover Page
        self._add_cover_page(analysis, "NSF")
        
        # Executive Summary
        self._add_section("EXECUTIVE SUMMARY", proposal.get('executive_summary', ''))
        
        # Project Description
        self._add_section("PROJECT DESCRIPTION", proposal.get('project_description', ''))
        
        # Research Plan
        self._add_section("RESEARCH PLAN", proposal.get('research_plan', ''))
        
        # Broader Impacts
        self._add_section("BROADER IMPACTS", proposal.get('broader_impacts', ''))
        
        # Budget Justification
        self._add_section("BUDGET JUSTIFICATION", proposal.get('budget_justification', ''))
        
        # Timeline
        self._add_timeline_section(proposal.get('timeline', {}))
        
        # References
        self._add_section("REFERENCES", proposal.get('references', ''))
        
        # Appendix: Analysis Summary
        self.doc.add_page_break()
        self._add_appendix(analysis, evaluation, innovations)


    def _format_nih_proposal(self, data: Dict):
        """Format as NIH-style proposal"""
        # Similar to NSF but with NIH-specific sections
        self._format_nsf_proposal(data)  # Reuse for now
    
    def _format_simple_proposal(self, data: Dict):
        """Format as simple proposal"""
        
        analysis = data.get('analysis', {})
        proposal = data.get('proposal', {}).get('proposal', {})
        
        # Title
        title = self.doc.add_paragraph(analysis.get('title', 'Research Proposal'))
        title.style = 'CustomTitle'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # All sections
        for section_name, section_content in proposal.items():
            if section_name == 'timeline':
                self._add_timeline_section(section_content)
            else:
                self._add_section(section_name.replace('_', ' ').upper(), str(section_content))
    
    def _add_cover_page(self, analysis: Dict, agency: str):
        """Add cover page"""
        
        # Title
        title = self.doc.add_paragraph(f"GRANT PROPOSAL\n{agency}")
        title.style = 'CustomTitle'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Project title
        project_title = self.doc.add_heading(
            f"Extension and Application of:\n{analysis.get('title', 'Unknown')}", 
            level=2
        )
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Info table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('Principal Investigator:', '[PI Name]'),
            ('Institution:', '[Institution Name]'),
            ('Duration:', '3 years'),
            ('Requested Amount:', '$500,000'),
            ('Submission Date:', datetime.now().strftime('%B %d, %Y')),
            ('Based on:', ', '.join(analysis.get('authors', ['Unknown'])[:3]))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        self.doc.add_page_break()
    
    def _add_section(self, title: str, content: str):
        """Add a section to document"""
        
        # Section heading
        self.doc.add_heading(title, level=1)
        
        # Content (split by paragraphs)
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                self.doc.add_paragraph(para_text.strip())
        
        self.doc.add_paragraph()  # Spacing

    def _add_timeline_section(self, timeline: Dict):
        """Add timeline section with table"""
        
        self.doc.add_heading("PROJECT TIMELINE", level=1)
        
        if not timeline:
            self.doc.add_paragraph("Timeline to be determined.")
            return
        
        # Create table
        num_years = len(timeline)
        table = self.doc.add_table(rows=num_years + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.rows[0].cells[0].text = "Period"
        table.rows[0].cells[1].text = "Milestones"
        
        # Make header bold
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data
        for i, (year, milestones) in enumerate(timeline.items(), 1):
            table.rows[i].cells[0].text = year
            
            if isinstance(milestones, list):
                milestone_text = '\n'.join(f"• {m}" for m in milestones)
            else:
                milestone_text = str(milestones)
            
            table.rows[i].cells[1].text = milestone_text
        
        self.doc.add_paragraph()
    

    def _add_appendix(self, analysis: Dict, evaluation: Dict, innovations: Dict):
        """Add appendix with analysis details"""
        
        self.doc.add_heading("APPENDIX: DETAILED ANALYSIS", level=1)
        
        # Paper info
        self.doc.add_heading("Paper Information", level=2)
        info_para = self.doc.add_paragraph()
        info_para.add_run(f"Title: ").bold = True
        info_para.add_run(f"{analysis.get('title', 'Unknown')}\n")
        info_para.add_run(f"Authors: ").bold = True
        info_para.add_run(f"{', '.join(analysis.get('authors', ['Unknown'])[:5])}\n")
        if analysis.get('year'):
            info_para.add_run(f"Year: ").bold = True
            info_para.add_run(f"{analysis['year']}\n")
        
        # Quality scores
        self.doc.add_heading("Quality Assessment Scores", level=2)
        
        scores = evaluation.get('scores', {})
        table = self.doc.add_table(rows=len(scores) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Score"
        
        for i, (metric, score) in enumerate(scores.items(), 1):
            table.rows[i].cells[0].text = metric.capitalize()
            table.rows[i].cells[1].text = f"{score}/10"
        
        self.doc.add_paragraph()
        
        # Key contributions
        self.doc.add_heading("Key Contributions", level=2)
        for contrib in analysis.get('key_contributions', [])[:5]:
            self.doc.add_paragraph(contrib, style='List Bullet')
        
        # Future directions
        self.doc.add_heading("Future Research Directions", level=2)
        for direction in innovations.get('future_directions', [])[:5]:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(direction.get('direction', 'N/A')).bold = True
            para.add_run(f"\n{direction.get('description', '')}")
    
    def export_to_pdf(
        self,
        proposal_data: Dict,
        output_path: str = "grant_proposal.pdf",
        method: str = "docx2pdf"
    ) -> str:
        """
        Export to PDF format
        
        Args:
            proposal_data: Complete proposal data
            output_path: Path for output file
            method: 'docx2pdf' or 'reportlab'
        
        Returns:
            Path to created file
        """
        print(f"📄 Exporting to PDF: {output_path}")
        
        if method == "docx2pdf":
            # First create DOCX
            docx_path = output_path.replace('.pdf', '.docx')
            self.export_to_docx(proposal_data, docx_path)
            
            # Convert to PDF
            try:
                from docx2pdf import convert
                convert(docx_path, output_path)
                print(f"✅ PDF saved: {output_path}")
                return output_path
            except ImportError:
                print("⚠️  docx2pdf not installed. Install with: pip install docx2pdf")
                print("   Returning DOCX file instead.")
                return docx_path
        
        elif method == "reportlab":
            # Use reportlab for direct PDF generation
            return self._export_pdf_reportlab(proposal_data, output_path)
        
        else:
            raise ValueError(f"Unknown method: {method}")
    
    def _export_pdf_reportlab(self, proposal_data: Dict, output_path: str) -> str:
        """Export using reportlab (direct PDF)"""
        
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            
            # Create PDF
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='navy',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            # Get data
            analysis = proposal_data.get('analysis', {})
            proposal = proposal_data.get('proposal', {}).get('proposal', {})
            
            # Title
            story.append(Paragraph("GRANT PROPOSAL", title_style))
            story.append(Paragraph(analysis.get('title', 'Unknown'), styles['Heading2']))
            story.append(Spacer(1, 0.5*inch))
            
            # Sections
            for section_name, section_content in proposal.items():
                if section_name != 'timeline':
                    story.append(Paragraph(section_name.replace('_', ' ').upper(), styles['Heading1']))
                    
                    # Split into paragraphs
                    paragraphs = str(section_content).split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            story.append(Paragraph(para_text.strip(), styles['BodyText']))
                            story.append(Spacer(1, 0.2*inch))
                    
                    story.append(Spacer(1, 0.3*inch))
            
            # Build PDF
            doc.build(story)
            print(f"✅ PDF saved: {output_path}")
            return output_path
            
        except ImportError:
            print("⚠️  reportlab not installed. Install with: pip install reportlab")
            return self.export_to_docx(proposal_data, output_path.replace('.pdf', '.docx'))
    
    def export_to_html(self, proposal_data: Dict, output_path: str = "grant_proposal.html") -> str:
        """Export to HTML format"""
        
        print(f"🌐 Exporting to HTML: {output_path}")
        
        analysis = proposal_data.get('analysis', {})
        evaluation = proposal_data.get('evaluation', {})
        innovations = proposal_data.get('innovations', {})
        proposal = proposal_data.get('proposal', {}).get('proposal', {})
        
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Grant Proposal - {analysis.get('title', 'Unknown')}</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1 {{
            color: #003366;
            border-bottom: 2px solid #003366;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #0066CC;
            margin-top: 30px;
        }}
        .cover {{
            text-align: center;
            margin-bottom: 50px;
            padding: 40px;
            background: #f5f5f5;
            border-radius: 10px;
        }}
        .section {{
            margin-bottom: 30px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            padding: 10px;
            border: 1px solid #ddd;
            text-align: left;
        }}
        th {{
            background-color: #003366;
            color: white;
        }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>GRANT PROPOSAL</h1>
        <h2>{analysis.get('title', 'Unknown')}</h2>
        <p><strong>Based on work by:</strong> {', '.join(analysis.get('authors', ['Unknown'])[:5])}</p>
        <p><strong>Requested Amount:</strong> $500,000 over 3 years</p>
    </div>
"""
        
        # Add sections
        for section_name, section_content in proposal.items():
            if section_name != 'timeline':
                html += f"""
    <div class="section">
        <h1>{section_name.replace('_', ' ').upper()}</h1>
        <p>{str(section_content).replace(chr(10), '</p><p>')}</p>
    </div>
"""
        
        # Timeline
        if 'timeline' in proposal:
            html += """
    <div class="section">
        <h1>PROJECT TIMELINE</h1>
        <table>
            <tr>
                <th>Period</th>
                <th>Milestones</th>
            </tr>
"""
            for year, milestones in proposal['timeline'].items():
                milestone_html = '<br>'.join(f"• {m}" for m in milestones) if isinstance(milestones, list) else str(milestones)
                html += f"""
            <tr>
                <td><strong>{year}</strong></td>
                <td>{milestone_html}</td>
            </tr>
"""
            html += """
        </table>
    </div>
"""
        
        html += """
</body>
</html>
"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        print(f"✅ HTML saved: {output_path}")
        return output_path
    
    def export_to_markdown(self, proposal_data: Dict, output_path: str = "grant_proposal.md") -> str:
        """Export to Markdown format"""
        
        print(f"📝 Exporting to Markdown: {output_path}")
        
        analysis = proposal_data.get('analysis', {})
        proposal = proposal_data.get('proposal', {}).get('proposal', {})
        
        md = f"""# Grant Proposal

## {analysis.get('title', 'Unknown')}

**Based on work by:** {', '.join(analysis.get('authors', ['Unknown'])[:5])}  
**Requested Amount:** $500,000 over 3 years  
**Date:** {datetime.now().strftime('%B %d, %Y')}

---

"""
        
        # Add sections
        for section_name, section_content in proposal.items():
            if section_name != 'timeline':
                md += f"## {section_name.replace('_', ' ').upper()}\n\n"
                md += f"{section_content}\n\n"
                md += "---\n\n"
        
        # Timeline
        if 'timeline' in proposal:
            md += "## PROJECT TIMELINE\n\n"
            for year, milestones in proposal['timeline'].items():
                md += f"### {year}\n\n"
                if isinstance(milestones, list):
                    for milestone in milestones:
                        md += f"- {milestone}\n"
                else:
                    md += f"{milestones}\n"
                md += "\n"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md)
        
        print(f"✅ Markdown saved: {output_path}")
        return output_path
    

# ==================== CLI ====================

def main():
    """Command-line interface"""
    import argparse
    import json
    
    parser = argparse.ArgumentParser(description="Export grant proposals to various formats")
    parser.add_argument("input_json", help="Input JSON file with proposal data")
    parser.add_argument("-f", "--format", choices=['docx', 'pdf', 'html', 'md', 'all'], default='docx', help="Output format")
    parser.add_argument("-o", "--output", help="Output filename (without extension)")
    parser.add_argument("-t", "--template", choices=['nsf', 'nih', 'simple'], default='nsf', help="Template style")
    
    args = parser.parse_args()
    
    # Load data
    with open(args.input_json, 'r') as f:
        data = json.load(f)
    
    # Determine output name
    if args.output:
        base_name = args.output
    else:
        base_name = "grant_proposal"
    
    # Create exporter
    exporter = ProposalExporter()
    
    # Export
    if args.format == 'all':
        exporter.export_to_docx(data, f"{base_name}.docx", args.template)
        exporter.export_to_pdf(data, f"{base_name}.pdf")
        exporter.export_to_html(data, f"{base_name}.html")
        exporter.export_to_markdown(data, f"{base_name}.md")
    elif args.format == 'docx':
        exporter.export_to_docx(data, f"{base_name}.docx", args.template)
    elif args.format == 'pdf':
        exporter.export_to_pdf(data, f"{base_name}.pdf")
    elif args.format == 'html':
        exporter.export_to_html(data, f"{base_name}.html")
    elif args.format == 'md':
        exporter.export_to_markdown(data, f"{base_name}.md")


if __name__ == "__main__":
    main()